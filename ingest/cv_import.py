"""Импорт резюме (pdf / docx / img / txt) → structured profile/cv.md.

Этот модуль отвечает за НАДЁЖНОЕ извлечение текста из файла резюме и за атомарную запись cv.md.
Структурирование «сырой текст → секции cv.md» в основном сценарии делает САМ Claude Code
по правилам из .claude/commands/onboard.md (главное — ничего не упускать).

Извлечение текста:
  - PDF  → pypdf (fallback pdfplumber для сложных макетов, если установлен).
  - DOCX → python-docx (параграфы + таблицы).
  - TXT/MD → как есть.
  - IMG  → в Claude Code читает зрение Claude; программный fallback — pytesseract (extra `ocr`).
"""

from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_CV_PATH = ROOT / "profile" / "cv.md"

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".text"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp"}


def extract_text(path: Path) -> str:
    """Извлечь сырой текст из резюме по расширению файла.

    Бросает FileNotFoundError, если файла нет; ValueError для неподдерживаемого формата;
    RuntimeError, если для изображения недоступен OCR (tesseract).
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Файл резюме не найден: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in IMAGE_SUFFIXES:
        return _extract_image(path)

    raise ValueError(
        f"Неподдерживаемый формат резюме: {suffix or '(без расширения)'}. "
        f"Поддерживаются: pdf, docx, txt/md, изображения ({', '.join(sorted(IMAGE_SUFFIXES))})."
    )


def _extract_docx(path: Path) -> str:
    """DOCX → текст: параграфы + ячейки таблиц (ничего не теряем)."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    """PDF → текст: pypdf; при пустом результате — fallback pdfplumber, если установлен."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(pages).strip()
    if text:
        return text

    # Fallback для сложных макетов/сканов с текстовым слоем.
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return text  # pdfplumber не установлен — отдаём что есть (возможно, скан без текста)
    with pdfplumber.open(str(path)) as pdf:
        return "\n".join((p.extract_text() or "") for p in pdf.pages).strip()


def _extract_image(path: Path) -> str:
    """IMG → текст через OCR (pytesseract). Требует системный бинарь `tesseract`.

    В Claude Code лучший путь — зрение Claude (изображение читается напрямую). Этот путь —
    программный fallback для автономного запуска.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise RuntimeError(
            "OCR недоступен: установите extra `ocr` (pip install -e '.[ocr]') "
            "и системный tesseract. В Claude Code прислать резюме картинкой можно — "
            "его прочитает зрение Claude; либо пришлите pdf/docx."
        ) from e

    try:
        return pytesseract.image_to_string(Image.open(str(path)), lang="rus+eng")
    except pytesseract.TesseractNotFoundError as e:
        raise RuntimeError(
            "Не найден бинарь tesseract (нужен для OCR изображений). "
            "Установите его (macOS: `brew install tesseract tesseract-lang`), "
            "или пришлите резюме в pdf/docx, или используйте Claude-зрение в /onboard."
        ) from e


def write_cv_markdown(markdown: str, dest: Path | None = None) -> Path:
    """Атомарно записать cv.md. Если файл уже есть — сохранить бэкап `<name>.bak`.

    Возвращает путь записанного файла.
    """
    dest = Path(dest) if dest is not None else DEFAULT_CV_PATH
    dest.parent.mkdir(parents=True, exist_ok=True)

    if dest.exists():
        backup = dest.with_name(dest.name + ".bak")
        backup.write_text(dest.read_text(encoding="utf-8"), encoding="utf-8")

    tmp = dest.with_name(dest.name + ".tmp")
    tmp.write_text(markdown, encoding="utf-8")
    tmp.replace(dest)  # атомарная замена
    return dest
