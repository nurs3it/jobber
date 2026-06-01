"""Тесты импорта резюме (Фаза 1): извлечение текста из pdf/docx/img/txt и запись cv.md.

Используем реальные файлы (round-trip), а не моки: создаём docx/pdf/png на лету.
"""

from __future__ import annotations

import shutil

import pytest

from ingest import cv_import


# --- extract_text: текстовые форматы ---

def test_extract_text_from_txt(tmp_path):
    p = tmp_path / "cv.txt"
    p.write_text("Иван Петров\nPython developer", encoding="utf-8")
    assert "Python developer" in cv_import.extract_text(p)


def test_extract_text_from_md(tmp_path):
    p = tmp_path / "cv.md"
    p.write_text("# Резюме\nReact, TypeScript", encoding="utf-8")
    assert "TypeScript" in cv_import.extract_text(p)


# --- extract_text: docx (параграфы + таблицы) ---

def test_extract_text_from_docx(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Иван Иванов")
    doc.add_paragraph("Senior Frontend — React, TypeScript")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Навык"
    table.rows[0].cells[1].text = "Vue"  # данные из таблицы не должны теряться
    p = tmp_path / "cv.docx"
    doc.save(p)

    text = cv_import.extract_text(p)
    assert "Иван Иванов" in text
    assert "React" in text
    assert "Vue" in text


# --- extract_text: pdf ---

def test_extract_text_from_pdf(tmp_path):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=14)
    pdf.cell(0, 10, "Frontend Developer CV")
    p = tmp_path / "cv.pdf"
    pdf.output(str(p))

    text = cv_import.extract_text(p)
    assert "Frontend Developer CV" in text


# --- extract_text: ошибки ---

def test_extract_text_unknown_extension_raises(tmp_path):
    p = tmp_path / "cv.xyz"
    p.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError):
        cv_import.extract_text(p)


def test_extract_text_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        cv_import.extract_text(tmp_path / "nope.pdf")


# --- extract_text: изображение (OCR fallback) ---

def test_extract_text_from_image_ocr(tmp_path):
    if shutil.which("tesseract") is None:
        pytest.skip("tesseract не установлен — OCR happy path пропущен")
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (500, 120), "white")
    ImageDraw.Draw(img).text((10, 50), "Backend Engineer", fill="black")
    p = tmp_path / "cv.png"
    img.save(p)

    text = cv_import.extract_text(p).lower()
    assert "backend" in text or "engineer" in text


def test_extract_text_image_without_tesseract_raises_clear_error(tmp_path):
    if shutil.which("tesseract") is not None:
        pytest.skip("tesseract установлен — happy path проверяется отдельно")
    from PIL import Image

    p = tmp_path / "cv.png"
    Image.new("RGB", (50, 50), "white").save(p)
    with pytest.raises(RuntimeError):
        cv_import.extract_text(p)


# --- write_cv_markdown ---

def test_write_cv_markdown_creates_file(tmp_path):
    dest = tmp_path / "profile" / "cv.md"
    out = cv_import.write_cv_markdown("# CV\nhello", dest)
    assert out == dest
    assert dest.read_text(encoding="utf-8") == "# CV\nhello"


def test_write_cv_markdown_backs_up_existing(tmp_path):
    dest = tmp_path / "cv.md"
    dest.write_text("OLD", encoding="utf-8")
    cv_import.write_cv_markdown("NEW", dest)
    assert dest.read_text(encoding="utf-8") == "NEW"
    bak = dest.with_name(dest.name + ".bak")
    assert bak.read_text(encoding="utf-8") == "OLD"
